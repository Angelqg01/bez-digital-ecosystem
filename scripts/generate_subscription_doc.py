from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading('BeZhas Hub - Planes de Suscripción', level=1)

doc.add_paragraph('Origen de la información:')
doc.add_paragraph(' - App: App\'s secundarias/Bezhas-Hub', style='List Bullet')
doc.add_paragraph(' - Documentos analizados: Conexión API-Hub — TAREAS.md, Conexión API-Hub .md, BEZHAS_CONEXION_TERCEROS_OPENCLAW_AEGIS.txt, bezhas-pay-system.jsx', style='List Bullet')

doc.add_heading('Resumen de planes de suscripción', level=2)
doc.add_paragraph('Los planes identificados en los documentos y la lógica de suscripción son los siguientes:')

plans = [
    {
        'name': 'STARTER',
        'price': '9,99 €/mes',
        'details': [
            '1 plataforma a elegir',
            'Rate limit: 1.000 req/día',
            'Webhooks: 5 eventos/hora',
            'Soporte: Chat OpenCLaw básico',
        ],
    },
    {
        'name': 'PRO',
        'price': '29,99 €/mes',
        'details': [
            'Hasta 4 plataformas',
            'Rate limit: 10.000 req/día',
            'Webhooks: ilimitados',
            'Soporte: OpenCLaw avanzado + AEGIS monitoring',
            'Sync automático cada 15 min',
        ],
    },
    {
        'name': 'ENTERPRISE',
        'price': '99,99 €/mes',
        'details': [
            'TODAS las plataformas (12)',
            'Rate limit: 100.000 req/día',
            'Webhooks: ilimitados + priority queue',
            'Soporte: OpenCLaw dedicado + AEGIS full autónomo',
            'Sync en tiempo real',
            'Custom adapters bajo solicitud',
            'SLA 99.9%',
        ],
    },
    {
        'name': 'VIP/WHALE',
        'price': 'Custom',
        'details': [
            'Todo Enterprise +',
            'Nodo AEGIS dedicado',
            'Agente OpenCLaw privado',
            'Integraciones custom ilimitadas',
            'White-label disponible',
            'API priority routing',
        ],
    },
]

for plan in plans:
    doc.add_heading(f"{plan['name']} - {plan['price']}", level=3)
    for detail in plan['details']:
        doc.add_paragraph(detail, style='List Bullet')


doc.add_heading('Planes de suscripción en la UI de la app', level=2)
doc.add_paragraph('La interfaz de suscripción detectada en bezhas-pay-system.jsx define planes y experiencia de compra:')

doc.add_paragraph('Planes implementados en la UI:', style='List Bullet')
for ui_plan in ['Free', 'Starter', 'Pro', 'Enterprise']:
    doc.add_paragraph(ui_plan, style='List Bullet 2')

doc.add_paragraph('Características de la experiencia de suscripción:')
for item in [
    'Selección visual de planes con icono, precio y lista de features.',
    'Pago con BEZ o criptos: BEZ, USDT, USDC, BNB, ETH.',
    'Descuento 20% al pagar con BEZ-Coin.',
    'Auto-renovación mensual mediante Smart Contract en BNB Chain.',
    'Botón de suscripción con texto dinámico según el plan.',
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('Flujo de suscripción previsto', level=2)
doc.add_paragraph('El plan documentado describe un onboarding automático con estos pasos:')
for item in [
    'Cliente realiza POST a /api/subscription/create.',
    'Payload incluye plan, método de pago, walletAddress y plataformas.',
    'Stripe/BEZ-Pay confirma la compra.',
    'Webhook dispara evento subscription.created.',
    'OpenClaw provisiona credenciales automáticamente.',
]:
    doc.add_paragraph(item, style='List Bullet')


doc.add_heading('Endpoints relacionados', level=2)
for endpoint in [
    'GET /subscription/plans',
    'POST /vip/subscribe',
    'POST /vip/upgrade',
    'POST /api/subscription/create',
]:
    doc.add_paragraph(endpoint, style='List Bullet')


doc.add_heading('Notas de arquitectura y contexto', level=2)
doc.add_paragraph('Esta app describe una conexión API-Hub donde el backend real en backend/ es el origen de verdad y las nuevas funciones deben montarse sobre una capa versionada y testeada.', style='List Bullet')
for note in [
    'API keys por plan/tier',
    'Rate-limit por tier',
    'Pagos reales con BEZ',
    'Orquestación de OpenClaw + AEGIS',
    'Datos reales en lugar de mocks',
]:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph('Se identificaron dos áreas de suscripción: la oferta de plan de productos/servicios con precios en euros y la experiencia de compra interna con precios en BEZ/USD.', style='Normal')

output_path = r"d:\BeZhas-Blockchain\App's secundarias\Bezhas-Hub\Conexión API-Hub - Planes de Suscripción.docx"
doc.save(output_path)
print(f"Documento generado en: {output_path}")
